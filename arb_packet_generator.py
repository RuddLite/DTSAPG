"""
arb_packet_generator.py  –  v1.0-TEST

STANDALONE TEST BUILD for Render.

This variant assumes the demand-letter-generator program does not exist:
every loss-of-use / rate value that the production version calculates from
demand_letter_generator.py's state+vehicle rate tables (REPL_DAILY,
DRIVER_WAGE_HOURLY, DRIVER_WAGE_DAILY, FINAL_DAILY_RATE, RAW_DAYS,
TOTAL_DAYS, LOSS_OF_USE_AMOUNT) — plus DIMINISHED_VALUE, which the
production version derives as a flat 10% of property damage — are instead
typed directly by the agent on the dashboard. There is no dependency on
demand_letter_generator.py in this build.

Generates a two-document Arbitration Packet (TOP.docx + BOTTOM.docx) from
agent-supplied intake data. The TOP and BOTTOM arbitration templates are a
SINGLE master template each, containing:

    1. Plain tokens          ->  [SOME_TOKEN]                (always replaced)
    2. Conditional clauses   ->  [IF KEY = YES]: sentence...  (kept/removed)
                                 [IF KEY1 = YES, KEY2 = NO]: sentence...
                                 [IF BARE_KEY]: sentence...   (bare == "YES")
    3. Optional bullet lines ->  [TOKEN] IF NONE OR LEFT BLANK, REMOVE LINE
    4. State case-law inserts -> [INSERT STATE_OF_ACCIDENT_CASE_LAW HERE]
                                 [INSERT_STATE_OF_ACCIDENT_PROGRESSIVE_CASE_LAW HERE]

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.text.paragraph import Paragraph

# ─────────────────────────────── PATHS ──────────────────────────────────────
BASE_DIR       = Path(__file__).parent
TOP_TEMPLATE   = BASE_DIR / "templates" / "TOP_Template.docx"
BOTTOM_TEMPLATE = BASE_DIR / "templates" / "BOTTOM_Template.docx"
CASE_LAW_DOC   = BASE_DIR / "templates" / "Case_Law_for_BOTTOM.docx"


STATE_ABBR = [
    "AL", "AK", "AZ", "AR", "CA", "CO", "CT", "DE", "FL", "GA", "HI", "ID",
    "IL", "IN", "IA", "KS", "KY", "LA", "ME", "MD", "MA", "MI", "MN", "MS",
    "MO", "MT", "NE", "NV", "NH", "NJ", "NM", "NY", "NC", "ND", "OH", "OK",
    "OR", "PA", "RI", "SC", "SD", "TN", "TX", "UT", "VT", "VA", "WA", "WV",
    "WI", "WY", "DC",
]
STATE_NAMES: Dict[str, str] = {
    "AL": "Alabama", "AK": "Alaska", "AZ": "Arizona", "AR": "Arkansas",
    "CA": "California", "CO": "Colorado", "CT": "Connecticut",
    "DE": "Delaware", "FL": "Florida", "GA": "Georgia", "HI": "Hawaii",
    "ID": "Idaho", "IL": "Illinois", "IN": "Indiana", "IA": "Iowa",
    "KS": "Kansas", "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine",
    "MD": "Maryland", "MA": "Massachusetts", "MI": "Michigan",
    "MN": "Minnesota", "MS": "Mississippi", "MO": "Missouri",
    "MT": "Montana", "NE": "Nebraska", "NV": "Nevada",
    "NH": "New Hampshire", "NJ": "New Jersey", "NM": "New Mexico",
    "NY": "New York", "NC": "North Carolina", "ND": "North Dakota",
    "OH": "Ohio", "OK": "Oklahoma", "OR": "Oregon", "PA": "Pennsylvania",
    "RI": "Rhode Island", "SC": "South Carolina", "SD": "South Dakota",
    "TN": "Tennessee", "TX": "Texas", "UT": "Utah", "VT": "Vermont",
    "VA": "Virginia", "WA": "Washington", "WV": "West Virginia",
    "WI": "Wisconsin", "WY": "Wyoming", "DC": "District of Columbia",
}

# States we currently have case law loaded for (grows as more are added to
# Case_Law_for_BOTTOM.docx — nothing else in the code needs to change).
def states_with_case_law() -> List[str]:
    return sorted(load_case_law().keys())


# ─────────────────────────── CASE LAW LOADER ─────────────────────────────────
_CASE_LAW_TAG_RE = re.compile(
    r"^\[([A-Z]{2})_(CASE_LAW|PROGRESSIVE_CASE_LAW)\]:?\s*(.*)$", re.I
)

_case_law_cache: Dict[str, Dict[str, List[str]]] | None = None


def load_case_law(path: Path = CASE_LAW_DOC) -> Dict[str, Dict[str, List[str]]]:
    """Parse Case_Law_for_BOTTOM.docx into:

        {"NY": {"case_law": [para, para, ...],
                "progressive_case_law": [para, ...]},
         "TN": {...}}

    Each state's [XX_CASE_LAW] / [XX_PROGRESSIVE_CASE_LAW] tag opens a
    section; every following non-blank paragraph belongs to that section
    until the next tag is hit. This lets a state's case-law block span
    multiple paragraphs (e.g. a lead-in sentence + one or more quotes),
    exactly as authored in the source document.
    """
    global _case_law_cache
    if _case_law_cache is not None:
        return _case_law_cache

    doc = Document(path)
    result: Dict[str, Dict[str, List[str]]] = {}
    current: Tuple[str, str] | None = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = _CASE_LAW_TAG_RE.match(text)
        if m:
            state, kind, remainder = m.group(1).upper(), m.group(2).lower(), m.group(3).strip()
            current = (state, kind)
            result.setdefault(state, {"case_law": [], "progressive_case_law": []})
            if remainder:
                result[state][kind].append(remainder)
        elif current:
            state, kind = current
            result[state][kind].append(text)

    _case_law_cache = result
    return result


# ───────────────────────── RUN / PARAGRAPH HELPERS ───────────────────────────
_BRACKET_RE = re.compile(r"\[[^\[\]]*\]")


def _merge_bracket_runs(paragraph: Paragraph) -> None:
    """Docx frequently splits a single [TOKEN] across several runs (e.g. one
    run per keystroke of spell-check/auto-correct history). Merge any runs
    that together form one complete [....] bracket into a single run so
    every downstream regex can assume a bracket is never split across runs.
    Formatting of the *first* run in the bracket is kept.
    """
    runs = paragraph.runs
    if not runs:
        return

    # map each character position in the full text to (run_index, local_pos)
    full_text = "".join(r.text for r in runs)
    for m in _BRACKET_RE.finditer(full_text):
        start, end = m.start(), m.end()
        # find run indices spanned by [start, end)
        pos = 0
        first_run_idx = None
        last_run_idx = None
        for i, r in enumerate(runs):
            r_start, r_end = pos, pos + len(r.text)
            if first_run_idx is None and r_start <= start < r_end:
                first_run_idx = i
            if r_start < end <= r_end:
                last_run_idx = i
                break
            pos = r_end
        if first_run_idx is None or last_run_idx is None or first_run_idx == last_run_idx:
            continue
        # merge text of [first_run_idx .. last_run_idx] into first_run_idx
        merged = "".join(runs[i].text for i in range(first_run_idx, last_run_idx + 1))
        runs[first_run_idx].text = merged
        for i in range(first_run_idx + 1, last_run_idx + 1):
            runs[i].text = ""


def _normalize_brackets(doc: Document) -> None:
    for p in doc.paragraphs:
        _merge_bracket_runs(p)


def _remove_paragraph(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)


def _strip_prefix_chars(paragraph: Paragraph, n: int) -> None:
    """Remove the first *n* characters of a paragraph's text, walking
    across run boundaries as needed (run formatting of remaining text is
    preserved)."""
    remaining = n
    for r in paragraph.runs:
        if remaining <= 0:
            break
        take = min(remaining, len(r.text))
        r.text = r.text[take:]
        remaining -= take


# ───────────────────────────── CONDITION LOGIC ───────────────────────────────
_IF_RE = re.compile(r"^\[IF\s+([^\]]+)\]:?\s*", re.I)


def _norm_key(s: str) -> str:
    """Collapse/strip all whitespace from a condition key so small typos in
    the source template (e.g. an accidental space: 'LANE_DIRECTION_
    AND_ROAD_KNOWN') still match the flags dict key built from the field id."""
    return re.sub(r"\s+", "", s).upper()


def _eval_condition(cond_str: str, flags: Dict[str, str]) -> bool:
    """cond_str examples:
        'FAVORABLE_PR = YES'
        'DRIVERS_STATEMENT = YES, SUPERVISOR_STATEMENT = NO'
        'ADVERSE_STATES_NO_LOU_SUPPORT_BUT_PAID_FULL_PD'   (bare -> == YES)

    Flag values are normally "YES"/"NO", but a field can also carry "NA"
    (agent selected N/A). Since a template condition only ever tests for an
    explicit "= YES" or "= NO", an "NA" flag value will not equal either,
    so the clause is correctly omitted either way without any special-case
    logic here.
    """
    for part in cond_str.split(","):
        part = part.strip()
        if not part:
            continue
        if "=" in part:
            key, val = part.split("=", 1)
            key, val = _norm_key(key), val.strip().upper()
        else:
            key, val = _norm_key(part), "YES"
        if flags.get(key, "NO").strip().upper() != val:
            return False
    return True


def _process_conditionals(doc: Document, flags: Dict[str, str]) -> None:
    """Pass 1: remove/keep every [IF ...]: clause paragraph."""
    for p in list(doc.paragraphs):
        full_text = "".join(r.text for r in p.runs)
        m = _IF_RE.match(full_text)
        if not m:
            continue
        keep = _eval_condition(m.group(1), flags)
        if not keep:
            _remove_paragraph(p)
            continue
        _strip_prefix_chars(p, m.end())
        # if nothing meaningful remains after stripping the tag, drop it
        if not "".join(r.text for r in p.runs).strip():
            _remove_paragraph(p)


# ───────────────────── OPTIONAL BULLET-LINE FILTERING ────────────────────────
_BLANK_LINE_MARKER = re.compile(
    r"\s*IF NONE OR LEFT BLANK,\s*REMOVE LINE", re.I
)
_LINE_TOKEN_RE = re.compile(r"\[([^\[\]]+)\]")


def _process_optional_bullets(doc: Document, ctx: Dict[str, str]) -> None:
    """Some paragraphs (e.g. the NEGLIGENCE THEORY bullet list) contain
    multiple lines separated by literal '\\n' characters inside run text.
    Any line ending in 'IF NONE OR LEFT BLANK, REMOVE LINE' should be
    dropped entirely if its bracketed token has no value; otherwise the
    marker text itself is stripped and the line is kept.
    """
    for p in list(doc.paragraphs):
        full_text = "".join(r.text for r in p.runs)
        if "REMOVE LINE" not in full_text.upper():
            continue

        lines = full_text.split("\n")
        keep_lines: List[bool] = []
        cleaned_lines: List[str] = []
        for line in lines:
            marker_match = _BLANK_LINE_MARKER.search(line)
            if marker_match:
                tok_match = _LINE_TOKEN_RE.search(line)
                tok_key = tok_match.group(1).strip() if tok_match else None
                value = (ctx.get(tok_key, "") if tok_key else "").strip()
                if not value:
                    keep_lines.append(False)
                    cleaned_lines.append(line)  # unused
                    continue
                line = line[: marker_match.start()] + line[marker_match.end():]
            keep_lines.append(True)
            cleaned_lines.append(line)

        new_text = "\n".join(l for l, k in zip(cleaned_lines, keep_lines) if k)

        # rewrite paragraph: keep formatting of first run, drop the rest,
        # then re-insert remaining lines with explicit line breaks.
        if not p.runs:
            continue
        first_run = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
        first_run.text = ""
        parts = new_text.split("\n")
        for i, part in enumerate(parts):
            if i > 0:
                first_run.add_break()
            first_run.add_text(part)


# ─────────────────────────── CASE LAW INSERTION ──────────────────────────────
_CASE_LAW_PLACEHOLDER_RE = re.compile(
    r"\[INSERT\s+STATE_OF_ACCIDENT_CASE_LAW\s+HERE\]:?\s*", re.I
)
_PROGRESSIVE_PLACEHOLDER_RE = re.compile(
    r"\[INSERT_?STATE_OF_ACCIDENT_PROGRESSIVE_CASE_LAW\s+HERE\]\s*", re.I
)


def _insert_state_case_law(doc: Document, state: str) -> None:
    """Replace the standalone `[INSERT STATE_OF_ACCIDENT_CASE_LAW HERE]:`
    paragraph with one paragraph per source case-law paragraph (so quotes
    keep their own line, matching the source formatting), and substitute
    the inline `[INSERT_STATE_OF_ACCIDENT_PROGRESSIVE_CASE_LAW HERE]` token
    with the progressive-carrier case law as a single inline block.
    """
    case_law = load_case_law()
    state_entry = case_law.get(state, {})
    case_paras = state_entry.get("case_law") or [
        f"[No case law currently on file for {STATE_NAMES.get(state, state)}. "
        "Please attach relevant case law manually.]"
    ]
    progressive_text = " ".join(state_entry.get("progressive_case_law") or [
        f"[No progressive-specific case law currently on file for "
        f"{STATE_NAMES.get(state, state)}.]"
    ])

    # standalone placeholder paragraph
    for p in list(doc.paragraphs):
        full_text = "".join(r.text for r in p.runs)
        if _CASE_LAW_PLACEHOLDER_RE.search(full_text):
            anchor = p._element
            for para_text in case_paras:
                new_p = doc.add_paragraph(para_text)
                # match body formatting of the template by copying pPr/style
                new_p.style = p.style
                anchor.addprevious(new_p._element)
            _remove_paragraph(p)
            break

    # inline progressive-case-law token (lives inside a larger paragraph)
    for p in doc.paragraphs:
        for r in p.runs:
            if _PROGRESSIVE_PLACEHOLDER_RE.search(r.text):
                r.text = _PROGRESSIVE_PLACEHOLDER_RE.sub(progressive_text + " ", r.text)


# ───────────────────────────── TOKEN SUBSTITUTION ────────────────────────────
_TOKEN_RE = re.compile(r"\[([^\[\]]+)\]")


def _clear_placeholder_color(run) -> None:
    """Templates color the [BRACKETED] placeholder text red so it stands
    out to whoever is filling in the form. Once real data replaces it,
    reset to the document's normal (automatic/black) text color."""
    from docx.oxml.ns import qn
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return
    color_el = rPr.find(qn("w:color"))
    if color_el is not None:
        rPr.remove(color_el)


def _substitute_tokens(doc: Document, ctx: Dict[str, str]) -> None:
    def _sub(run) -> None:
        if "[" not in run.text:
            return
        matched = False

        def repl(m: re.Match) -> str:
            nonlocal matched
            key = m.group(1).strip()
            if key in ctx:
                matched = True
                return ctx[key]
            return m.group(0)

        run.text = _TOKEN_RE.sub(repl, run.text)
        if matched:
            _clear_placeholder_color(run)

    for p in doc.paragraphs:
        for r in p.runs:
            _sub(r)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        _sub(r)


# ─────────────────────────────── MERGE PIPELINE ──────────────────────────────
def _merge(template: Path, dest: Path, ctx: Dict[str, str], flags: Dict[str, str],
           state: str | None = None) -> None:
    doc = Document(template)
    _normalize_brackets(doc)
    _process_conditionals(doc, flags)
    _process_optional_bullets(doc, ctx)
    if state:
        _insert_state_case_law(doc, state)
    _substitute_tokens(doc, ctx)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)


# ────────────────────────────────── CONTEXT BUILDER ──────────────────────────────

def build_context(data: Dict[str, Any]) -> Tuple[Dict[str, str], Dict[str, str]]:
    """Turn raw intake `data` (as submitted by the dashboard) into:
        ctx   – token -> replacement text
        flags – condition-key -> "YES"/"NO"   (drives clause keep/remove)
    """
    yn = lambda v: "YES" if str(v).strip().upper() in ("YES", "Y", "TRUE", "1", "ON") else "NO"
    yn3 = lambda v: str(v).strip().upper() if str(v).strip().upper() in ("YES", "NO", "NA") else "NO"
    money = lambda v: f"{float(v or 0):,.2f}"

    flags: Dict[str, str] = {
        "LANE_DIRECTION_AND_ROAD_KNOWN": yn(data.get("lane_direction_and_road_known")),
        "POLICE_REPORT": yn(data.get("police_report")),
        "VIDEO": yn(data.get("video")),
        "DRIVERS_STATEMENT": yn(data.get("drivers_statement")),
        "SUPERVISOR_STATEMENT": yn(data.get("supervisor_statement")),
        "MONITOR_STATEMENT": yn(data.get("monitor_statement")),
        "CONTRIBUTING_FACTOR_ON_PR_ADVERSE": yn3(data.get("contributing_factor_on_pr_adverse")),
        "CONTRIBUTING_FACTOR_ON_PR_CLIENT": yn3(data.get("contributing_factor_on_pr_client")),
        "ADVERSE_ISSUED_CITATION": yn3(data.get("adverse_issued_citation")),
        "PHOTOS_OF_DAMAGES": yn(data.get("photos_of_damages")),
        "LIABILITY_OBVIOUS": yn(data.get("liability_obvious")),
        "FAVORABLE_PR": yn(data.get("favorable_pr")),
        "TOTAL_LOSS_TAXI": yn(data.get("total_loss_taxi")),
        "ADVERSE_PAID_100%_PD": yn(data.get("adverse_paid_100_pct_pd")),
        "PD_AMOUNT_AGREED_UPON": yn(data.get("pd_amount_agreed_upon")),
        "ADVERSE_PAID_PARTIAL_PD": yn(data.get("adverse_paid_partial_pd")),
        "LIABILITY_CONTESTED": yn(data.get("liability_contested")),
        "PARTIAL_LOU": yn(data.get("partial_lou")),
        "PARTIAL_PAYMENT_COMBO_CHECK": yn(data.get("partial_payment_combo_check")),
        "ADVERSE_PAID_PARTIAL_LOU_WITHOUT_SUPPORTS": yn(data.get("adverse_paid_partial_lou_without_supports")),
        "ADVERSE_OFFERED_PARTIAL_LOU_WITHOUT_SUPPORTS": yn(data.get("adverse_offered_partial_lou_without_supports")),
        "PREVIOUS_VENDOR_LOU_DEMAND": yn(data.get("previous_vendor_lou_demand")),
        "PURSUING_DIMINISHED_VALUE": yn(data.get("pursuing_diminished_value")),
        "ADVERSE_STATES_NO_LOU_SUPPORT_BUT_PAID_FULL_PD": yn(data.get("adverse_states_no_lou_support_but_paid_full_pd")),
        "DRIVER_SALARY_REDUCTION": yn(data.get("driver_salary_reduction")),
        "ATTACHING_PREVIOUS_LOU_ARB_DECISIONS": yn(data.get("attaching_previous_lou_arb_decisions")),
        "ADVERSE_INSURANCE_CARRIER_IS_PROGRESSIVE": yn(data.get("adverse_insurance_carrier_is_progressive")),
    }

    state = str(data.get("accident_state", "")).upper().strip()

    property_damage = float(data.get("property_damage") or 0)
    towing = float(data.get("towing") or 0)
    loss_of_use_amount = float(data.get("loss_of_use_amount") or 0)
    dv_amt = float(data.get("diminished_value_amount") or 0) if flags["PURSUING_DIMINISHED_VALUE"] == "YES" else 0.0
    grand_total = property_damage + loss_of_use_amount + towing + dv_amt

    ctx: Dict[str, str] = {
        # ── narrative / TOP (Preamble) fields ──
        "ACCIDENT_DESCRIPTION": data.get("accident_description", ""),
        "LANE_OF_TRAVEL": data.get("lane_of_travel", ""),
        "DIRECTION_OF_TRAVEL": data.get("direction_of_travel", ""),
        "CLIENT_ROAD_OF_TRAVEL": data.get("client_road_of_travel", ""),
        "ADVERSE_LANE_OF_TRAVEL": data.get("adverse_lane_of_travel", ""),
        "ADVERSE_DIRECTION_OF_TRAVEL": data.get("adverse_direction_of_travel", ""),
        "ADVERSE_ROAD_OF_TRAVEL": data.get("adverse_road_of_travel", ""),
        "COLLISION_OCCURRENCE": data.get("collision_occurrence", ""),
        "PAGE_OF_PR_NARRATIVE": data.get("page_of_pr_narrative", ""),
        "PAGE_OF_DIAGRAM": data.get("page_of_diagram", ""),

        "ADVERSE_CONTRIBUTING_ACTION_2": data.get("adverse_contributing_action_2", ""),
        "ADVERSE_CONTRIBUTING_ACTION_3": data.get("adverse_contributing_action_3", ""),
        "ADVERSE_CONTRIBUTING_ACTION_4": data.get("adverse_contributing_action_4", ""),

        "PR_PAGE_CONTAINING_CONTRIBUTING_FACTORS_ADVERSE": data.get("pr_page_containing_contributing_factors_adverse", ""),
        "PR_CONTRIBUTING_ACTION_ADVERSE": data.get("pr_contributing_action_adverse", ""),
        "PR_PAGE_CONTAINING_CONTRIBUTING_FACTORS_CLIENT": data.get("pr_page_containing_contributing_factors_client", ""),

        "STATE_VTL": data.get("state_vtl", ""),
        "TRAFFIC_LAW": data.get("traffic_law", ""),
        "TRAFFIC_LAW_DESCRIPTION": data.get("traffic_law_description", ""),

        "PR_PAGE_SHOWING_ADVERSE_POINT_OF_IMPACT": data.get("pr_page_showing_adverse_point_of_impact", ""),
        "ADVERSE_POINT_OF_IMPACT": data.get("adverse_point_of_impact", ""),
        "PR_PAGE_SHOWING_CLIENT_POINT_OF_IMPACT": data.get("pr_page_showing_client_point_of_impact", ""),
        "CLIENT_POINT_OF_IMPACT": data.get("client_point_of_impact", ""),

        "TAXI_EQUIPMENT_TRANSFER_COST": money(data.get("taxi_equipment_transfer_cost")) if flags["TOTAL_LOSS_TAXI"] == "YES" else "",

        "PROPERTY_DAMAGE_AMOUNT": money(property_damage),
        "PROPERTY_DAMAGE_AMOUNT_AGREED_UPON": money(data.get("property_damage_amount_agreed_upon")),
        "PARTIAL_PD_AMOUNT": money(data.get("partial_pd_amount")),
        "PERCENTAGE_OF_ADVERSE_ESTIMATE_RECOVERED": str(data.get("percentage_of_adverse_estimate_recovered", "")).strip(),

        "LOSS_OF_USE_AMOUNT": money(loss_of_use_amount),
        "TOTAL_LOU_DAYS": str(data.get("total_lou_days", "")).strip(),
        "PARTIAL_LOSS_OF_USE_PAYMENT": money(data.get("partial_loss_of_use_payment")),
        "TOTAL_ADVERSE_PARTIAL": money(data.get("total_adverse_partial")),
        "TOTAL_ADVERSE_PD_PAYMENT": money(data.get("total_adverse_pd_payment")),

        "YEAR_PREVIOUS_DEMAND_SUBMITTED": str(data.get("year_previous_demand_submitted", "")),
        "YEAR_DTS_DEMAND_SENT": str(data.get("year_dts_demand_sent", "")),

        "DIMINISHED_VALUE": money(dv_amt) if flags["PURSUING_DIMINISHED_VALUE"] == "YES" else "",

        # ── BOTTOM (Contentions) / LOU fields (typed directly by the agent in this test build) ──
        "STATE_OF_ACCIDENT": STATE_NAMES.get(state, state),
        "REPL_DAILY": money(data.get("repl_daily")),
        "DRIVER_WAGE_HOURLY": money(data.get("driver_wage_hourly")),
        "DRIVER_WAGE_DAILY": money(data.get("driver_wage_daily")),
        "FINAL_DAILY_RATE": money(data.get("final_daily_rate")),
        "RAW_DAYS": str(data.get("raw_days", "")).strip(),
        "APPLICABLE_WEEKEND_DAYS": str(data.get("applicable_weekend_days", "")).strip(),
        "TOTAL_DAYS": str(data.get("total_lou_days", "")).strip(),
        "REPAIR_HOURS": str(int(float(data.get("repair_hours") or 0))),

        "GRAND_TOTAL": f"{grand_total:,.2f}",
    }

    return ctx, flags


# ─────────────────────────────── PUBLIC ENTRYPOINT ───────────────────────────
def make_arb_packet(data: Dict[str, Any], out_dir: str | Path = "output",
                     file_prefix: str = "arb_packet") -> Tuple[Path, Path]:
    """Build both documents. Returns (top_path, bottom_path)."""
    ctx, flags = build_context(data)
    state = str(data.get("accident_state", "")).upper().strip()

    out_dir = Path(out_dir)
    top_path = out_dir / f"{file_prefix}_PREAMBLE.docx"
    bottom_path = out_dir / f"{file_prefix}_CONTENTIONS.docx"

    _merge(TOP_TEMPLATE, top_path, ctx, flags, state=None)
    _merge(BOTTOM_TEMPLATE, bottom_path, ctx, flags, state=state)

    return top_path, bottom_path


# ────────────────────────────────── CLI ──────────────────────────────────────
def _cli():
    ap = argparse.ArgumentParser(description="Generate an arbitration packet (TOP + BOTTOM .docx).")
    ap.add_argument("--json", required=True, help="Path to a JSON file with the intake data")
    ap.add_argument("--out", default="output")
    args = ap.parse_args()

    import json
    data = json.loads(Path(args.json).read_text())
    top, bottom = make_arb_packet(data, out_dir=args.out)
    print("✅  TOP written    →", top)
    print("✅  BOTTOM written →", bottom)


if __name__ == "__main__":
    _cli()
